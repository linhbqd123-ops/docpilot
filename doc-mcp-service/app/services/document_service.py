"""Document processing service.

Handles all Word document operations: extraction, modification, and generation.
Uses python-docx internally but exposes a clean structured JSON interface.
"""

from __future__ import annotations

import base64
import io
import logging
import re
import uuid
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

from app.models.document import (
    BlockType,
    DocumentBlock,
    DocumentStructure,
    TextRun,
    RunFormat,
    TableRow,
    TableCell,
    BlockChange,
    DocumentResponse,
)

logger = logging.getLogger(__name__)


def _generate_block_id() -> str:
    return uuid.uuid4().hex[:8]


def _extract_run_format(run) -> Optional[RunFormat]:
    """Extract formatting from a python-docx Run object."""
    fmt = RunFormat()
    has_format = False

    if run.bold is not None:
        fmt.bold = run.bold
        has_format = True
    if run.italic is not None:
        fmt.italic = run.italic
        has_format = True
    if run.underline is not None:
        fmt.underline = bool(run.underline)
        has_format = True
    if run.font.name:
        fmt.font_name = run.font.name
        has_format = True
    if run.font.size:
        fmt.font_size = run.font.size.pt
        has_format = True
    if run.font.color and run.font.color.rgb:
        fmt.color = str(run.font.color.rgb)
        has_format = True

    return fmt if has_format else None


def _apply_run_format(run, fmt: RunFormat):
    """Apply RunFormat to a python-docx Run object."""
    if fmt is None:
        return
    if fmt.bold is not None:
        run.bold = fmt.bold
    if fmt.italic is not None:
        run.italic = fmt.italic
    if fmt.underline is not None:
        run.underline = fmt.underline
    if fmt.font_name:
        run.font.name = fmt.font_name
    if fmt.font_size:
        run.font.size = Pt(fmt.font_size)
    if fmt.color:
        try:
            run.font.color.rgb = RGBColor.from_string(fmt.color)
        except (ValueError, AttributeError):
            pass


def _is_list_paragraph(paragraph) -> bool:
    """Check if a paragraph is a list item by inspecting its XML."""
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is not None:
        numPr = pPr.find(qn('w:numPr'))
        return numPr is not None
    return False


def _get_list_level(paragraph) -> int:
    """Get the indentation level of a list item."""
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is not None:
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            ilvl = numPr.find(qn('w:ilvl'))
            if ilvl is not None:
                return int(ilvl.get(qn('w:val'), '0'))
    return 0


def _doc_from_base64(data: Optional[str]) -> Document:
    """Create a Document from base64-encoded data, or a new empty one."""
    if data:
        raw = base64.b64decode(data)
        return Document(io.BytesIO(raw))
    return Document()


def _doc_to_base64(doc: Document) -> str:
    """Serialize a Document to a base64-encoded string."""
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return base64.b64encode(buf.read()).decode("utf-8")


class DocumentService:
    """Encapsulates all document operations."""

    def extract_structure(
        self,
        document_base64: Optional[str] = None,
        include_formatting: bool = True,
    ) -> DocumentStructure:
        """Extract structured representation from a Word document."""
        doc = _doc_from_base64(document_base64)
        return self._extract_structure_from_doc(doc, include_formatting)

    def apply_changes(
        self,
        changes: list[BlockChange],
        document_base64: Optional[str] = None,
    ) -> DocumentResponse:
        """Apply a list of block-level changes to a document.

        Extracts structure from the already-loaded document object to avoid
        decoding base64 twice. Matches blocks by ID ordering and applies
        text replacements while preserving formatting.
        """
        doc = _doc_from_base64(document_base64)
        structure = self._extract_structure_from_doc(doc, include_formatting=True)

        # Build block_id -> index mapping
        id_to_index: dict[str, int] = {}
        for idx, block in enumerate(structure.blocks):
            id_to_index[block.id] = idx

        # Map paragraph index (order)
        para_blocks = [b for b in structure.blocks if b.type in (BlockType.PARAGRAPH, BlockType.LIST_ITEM)]
        para_id_to_doc_idx: dict[str, int] = {}
        for i, b in enumerate(para_blocks):
            para_id_to_doc_idx[b.id] = i

        # Apply changes
        paragraphs_to_delete: list[int] = []
        logger.debug("DocumentService.apply_changes start changes=%d", len(changes))

        for change in changes:
            if change.block_id not in id_to_index:
                logger.debug("Skipping unknown block_id=%s", change.block_id)
                continue

            block = structure.blocks[id_to_index[change.block_id]]

            if block.type in (BlockType.PARAGRAPH, BlockType.LIST_ITEM):
                doc_idx = para_id_to_doc_idx.get(change.block_id)
                if doc_idx is None or doc_idx >= len(doc.paragraphs):
                    continue

                para = doc.paragraphs[doc_idx]

                if change.action == "delete":
                    paragraphs_to_delete.append(doc_idx)
                    logger.debug("Marked paragraph index=%d for deletion", doc_idx)
                    continue

                if change.action == "update" and change.new_text is not None:
                    # Preserve formatting of first run, replace text
                    if change.new_runs:
                        # Clear existing runs and add new ones with formatting
                        for run in para.runs:
                            run.text = ""
                        for i, new_run in enumerate(change.new_runs):
                            if i < len(para.runs):
                                para.runs[i].text = new_run.text
                                if new_run.format:
                                    _apply_run_format(para.runs[i], new_run.format)
                            else:
                                r = para.add_run(new_run.text)
                                if new_run.format:
                                    _apply_run_format(r, new_run.format)
                    else:
                        # Simple text replacement - keep first run's format
                        if para.runs:
                            first_run_format = _extract_run_format(para.runs[0])
                            for run in para.runs:
                                run.text = ""
                            para.runs[0].text = change.new_text
                        else:
                            para.add_run(change.new_text)

                    if change.new_style:
                        try:
                            para.style = change.new_style
                            logger.debug(
                                "Applied new_style=%s to block_id=%s",
                                change.new_style,
                                change.block_id,
                            )
                        except KeyError:
                            logger.warning(
                                "Requested style not found: %s for block_id=%s",
                                change.new_style,
                                change.block_id,
                            )

                if change.action == "update" and change.new_style:
                    try:
                        para.style = change.new_style
                    except KeyError:
                        # Ignore invalid style names to keep apply_changes resilient.
                        pass

        # Delete marked paragraphs (reverse order to preserve indices)
        for idx in sorted(paragraphs_to_delete, reverse=True):
            if idx < len(doc.paragraphs):
                p_element = doc.paragraphs[idx]._element
                p_element.getparent().remove(p_element)

        new_base64 = _doc_to_base64(doc)
        new_structure = self.extract_structure(new_base64, include_formatting=True)
        logger.debug("DocumentService.apply_changes done resulting_blocks=%d", len(new_structure.blocks))

        return DocumentResponse(
            structure=new_structure,
            document_base64=new_base64,
            message=f"Applied {len(changes)} changes successfully",
        )

    def clear_document(self, document_base64: Optional[str] = None) -> DocumentResponse:
        """Clear all content from a document, preserving styles."""
        doc = _doc_from_base64(document_base64)

        # Remove all paragraphs except the first (required by python-docx)
        for i in range(len(doc.paragraphs) - 1, 0, -1):
            p_element = doc.paragraphs[i]._element
            p_element.getparent().remove(p_element)

        # Clear the first paragraph
        if doc.paragraphs:
            for run in doc.paragraphs[0].runs:
                run.text = ""
            doc.paragraphs[0].text = ""

        # Remove all tables
        for table in doc.tables:
            table._element.getparent().remove(table._element)

        new_base64 = _doc_to_base64(doc)
        return DocumentResponse(
            structure=DocumentStructure(blocks=[], metadata={}),
            document_base64=new_base64,
            message="Document cleared",
        )

    def insert_structured_content(
        self,
        blocks: list[DocumentBlock],
        position: str = "end",
        document_base64: Optional[str] = None,
    ) -> DocumentResponse:
        """Insert structured blocks into a document."""
        doc = _doc_from_base64(document_base64)
        logger.debug("DocumentService.insert_structured_content start blocks=%d position=%s", len(blocks), position)

        for block in blocks:
            if block.type == BlockType.TABLE:
                self._insert_table(doc, block)
            else:
                self._insert_paragraph(doc, block)

        new_base64 = _doc_to_base64(doc)
        new_structure = self.extract_structure(new_base64, include_formatting=True)
        logger.debug("DocumentService.insert_structured_content done resulting_blocks=%d", len(new_structure.blocks))

        return DocumentResponse(
            structure=new_structure,
            document_base64=new_base64,
            message=f"Inserted {len(blocks)} blocks",
        )

    def _extract_structure_from_doc(
        self,
        doc: Document,
        include_formatting: bool = True,
    ) -> DocumentStructure:
        """Extract structured representation directly from a Document object."""
        blocks: list[DocumentBlock] = []

        for paragraph in doc.paragraphs:
            block_id = _generate_block_id()
            style_name = paragraph.style.name if paragraph.style else None

            is_list = _is_list_paragraph(paragraph)
            block_type = BlockType.LIST_ITEM if is_list else BlockType.PARAGRAPH
            level = _get_list_level(paragraph) if is_list else None

            if style_name and style_name.startswith("Heading"):
                match = re.search(r"(\d+)", style_name)
                if match:
                    level = int(match.group(1))

            runs: list[TextRun] = []
            if include_formatting:
                for run in paragraph.runs:
                    fmt = _extract_run_format(run)
                    runs.append(TextRun(text=run.text, format=fmt))

            blocks.append(
                DocumentBlock(
                    id=block_id,
                    type=block_type,
                    style=style_name,
                    text=paragraph.text,
                    runs=runs,
                    level=level,
                )
            )

        for table in doc.tables:
            block_id = _generate_block_id()
            rows: list[TableRow] = []
            for row in table.rows:
                cells: list[TableCell] = []
                for cell in row.cells:
                    cell_runs: list[TextRun] = []
                    if include_formatting:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                fmt = _extract_run_format(run)
                                cell_runs.append(TextRun(text=run.text, format=fmt))
                    cells.append(
                        TableCell(
                            text=cell.text,
                            runs=cell_runs,
                            style=cell.paragraphs[0].style.name if cell.paragraphs else None,
                        )
                    )
                rows.append(TableRow(cells=cells))

            all_text = " | ".join(
                " | ".join(c.text for c in r.cells)
                for r in rows
            )
            blocks.append(
                DocumentBlock(
                    id=block_id,
                    type=BlockType.TABLE,
                    text=all_text[:200],
                    rows=rows,
                )
            )

        metadata = {}
        if doc.core_properties:
            cp = doc.core_properties
            if cp.title:
                metadata["title"] = cp.title
            if cp.author:
                metadata["author"] = cp.author

        return DocumentStructure(blocks=blocks, metadata=metadata)

    def _insert_paragraph(self, doc: Document, block: DocumentBlock):
        """Insert a paragraph block into the document."""
        para = doc.add_paragraph()

        # Apply style if specified
        if block.style:
            try:
                para.style = block.style
            except KeyError:
                pass

        if block.runs:
            for text_run in block.runs:
                run = para.add_run(text_run.text)
                if text_run.format:
                    _apply_run_format(run, text_run.format)
        else:
            para.add_run(block.text)

    def _insert_table(self, doc: Document, block: DocumentBlock):
        """Insert a table block into the document."""
        if not block.rows:
            return

        num_rows = len(block.rows)
        num_cols = max(len(row.cells) for row in block.rows) if block.rows else 0

        if num_cols == 0:
            return

        table = doc.add_table(rows=num_rows, cols=num_cols)
        
        # Try to apply Table Grid style, fall back to no style if not available
        try:
            table.style = "Table Grid"
            logger.debug("Applied table style Table Grid")
        except KeyError:
            # Table Grid style not available, use default table style
            try:
                table.style = "Table Normal"
                logger.debug("Applied fallback table style Table Normal")
            except KeyError:
                # No common table styles available, proceed without style
                logger.warning("No supported table style available (Table Grid/Table Normal). Using default style.")
                pass

        for r_idx, row in enumerate(block.rows):
            for c_idx, cell in enumerate(row.cells):
                if c_idx < num_cols:
                    doc_cell = table.cell(r_idx, c_idx)
                    doc_cell.text = cell.text


# Singleton instance
document_service = DocumentService()
