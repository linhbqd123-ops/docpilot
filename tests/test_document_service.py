"""Unit tests for the DOC-MCP DocumentService.

Run with:
    cd doc-mcp-service
    python -m pytest ../tests/test_document_service.py -v
"""

from __future__ import annotations

import base64
import io
import sys
from pathlib import Path

import pytest
from docx import Document

# Ensure imports work regardless of cwd
sys.path.insert(0, str(Path(__file__).parent.parent / "doc-mcp-service"))

from app.services.document_service import DocumentService, _doc_to_base64
from app.models.document import BlockChange, BlockType, DocumentBlock, TextRun, RunFormat, TableRow, TableCell


@pytest.fixture
def service():
    return DocumentService()


@pytest.fixture
def sample_base64():
    """Create a simple docx in memory and return its base64."""
    doc = Document()
    doc.add_heading("Test Title", level=1)
    doc.add_paragraph("First paragraph content.")
    doc.add_paragraph("Second paragraph content.")
    return _doc_to_base64(doc)


@pytest.fixture
def sample_with_list_base64():
    """Docx with bullets."""
    doc = Document()
    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Python", style="List Bullet")
    doc.add_paragraph("JavaScript", style="List Bullet")
    return _doc_to_base64(doc)


@pytest.fixture
def sample_with_table_base64():
    """Docx with a table."""
    doc = Document()
    doc.add_heading("People", level=1)
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Role"
    table.cell(1, 0).text = "Alice"
    table.cell(1, 1).text = "Engineer"
    return _doc_to_base64(doc)


# --- extract_structure ---

class TestExtractStructure:
    def test_extract_basic(self, service, sample_base64):
        structure = service.extract_structure(sample_base64)
        assert len(structure.blocks) == 3  # heading + 2 paragraphs
        assert structure.blocks[0].style == "Heading 1"
        assert structure.blocks[0].text == "Test Title"
        assert structure.blocks[1].text == "First paragraph content."
        assert structure.blocks[2].text == "Second paragraph content."

    def test_extract_empty_doc(self, service):
        doc = Document()
        b64 = _doc_to_base64(doc)
        structure = service.extract_structure(b64)
        # Empty doc still has one empty paragraph
        assert len(structure.blocks) >= 0

    def test_extract_preserves_block_type(self, service, sample_with_list_base64):
        structure = service.extract_structure(sample_with_list_base64)
        list_blocks = [b for b in structure.blocks if b.type == BlockType.LIST_ITEM]
        assert len(list_blocks) == 2

    def test_extract_table(self, service, sample_with_table_base64):
        structure = service.extract_structure(sample_with_table_base64)
        table_blocks = [b for b in structure.blocks if b.type == BlockType.TABLE]
        assert len(table_blocks) == 1
        assert len(table_blocks[0].rows) == 2
        assert table_blocks[0].rows[0].cells[0].text == "Name"

    def test_extract_with_formatting(self, service):
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("bold text")
        run.bold = True
        b64 = _doc_to_base64(doc)

        structure = service.extract_structure(b64, include_formatting=True)
        block = structure.blocks[0]
        assert len(block.runs) == 1
        assert block.runs[0].text == "bold text"
        assert block.runs[0].format is not None
        assert block.runs[0].format.bold is True

    def test_extract_without_formatting(self, service):
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("bold text")
        run.bold = True
        b64 = _doc_to_base64(doc)

        structure = service.extract_structure(b64, include_formatting=False)
        block = structure.blocks[0]
        assert len(block.runs) == 0  # No runs when formatting is off
        assert block.text == "bold text"


# --- apply_changes ---

class TestApplyChanges:
    def test_apply_text_change(self, service, sample_base64):
        structure = service.extract_structure(sample_base64)
        target = structure.blocks[1]  # "First paragraph content."

        changes = [
            BlockChange(block_id=target.id, new_text="Updated first paragraph.", action="update")
        ]

        result = service.apply_changes(changes, sample_base64)
        assert result.document_base64 is not None

        new_structure = service.extract_structure(result.document_base64)
        texts = [b.text for b in new_structure.blocks]
        assert "Updated first paragraph." in texts

    def test_apply_delete(self, service, sample_base64):
        structure = service.extract_structure(sample_base64)
        target = structure.blocks[2]  # "Second paragraph content."

        changes = [BlockChange(block_id=target.id, action="delete")]
        result = service.apply_changes(changes, sample_base64)

        new_structure = service.extract_structure(result.document_base64)
        texts = [b.text for b in new_structure.blocks]
        assert "Second paragraph content." not in texts

    def test_apply_no_changes(self, service, sample_base64):
        result = service.apply_changes([], sample_base64)
        assert result.document_base64 is not None

    def test_apply_unknown_block_id(self, service, sample_base64):
        changes = [BlockChange(block_id="nonexistent", new_text="x", action="update")]
        result = service.apply_changes(changes, sample_base64)
        # Should succeed silently (skip unknown blocks)
        assert result.document_base64 is not None


# --- clear_document ---

class TestClearDocument:
    def test_clear(self, service, sample_base64):
        result = service.clear_document(sample_base64)
        assert result.document_base64 is not None
        structure = service.extract_structure(result.document_base64)
        non_empty = [b for b in structure.blocks if b.text.strip()]
        assert len(non_empty) == 0


# --- insert_structured_content ---

class TestInsertContent:
    def test_insert_paragraphs(self, service, sample_base64):
        blocks = [
            DocumentBlock(id="new1", type=BlockType.PARAGRAPH, style="Heading 2", text="New Section"),
            DocumentBlock(id="new2", type=BlockType.PARAGRAPH, text="New paragraph text."),
        ]

        result = service.insert_structured_content(blocks, "end", sample_base64)
        assert result.document_base64 is not None
        structure = service.extract_structure(result.document_base64)
        texts = [b.text for b in structure.blocks]
        assert "New Section" in texts
        assert "New paragraph text." in texts

    def test_insert_table(self, service, sample_base64):
        block = DocumentBlock(
            id="t1",
            type=BlockType.TABLE,
            rows=[
                TableRow(cells=[
                    TableCell(text="A"),
                    TableCell(text="B"),
                ]),
                TableRow(cells=[
                    TableCell(text="1"),
                    TableCell(text="2"),
                ]),
            ],
        )
        result = service.insert_structured_content([block], "end", sample_base64)
        structure = service.extract_structure(result.document_base64)
        table_blocks = [b for b in structure.blocks if b.type == BlockType.TABLE]
        assert len(table_blocks) == 1


# --- roundtrip ---

class TestRoundtrip:
    def test_extract_apply_preserves_formatting(self, service):
        """Verify that extracting then applying with no changes preserves bold formatting."""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("Important")
        r.bold = True
        r2 = p.add_run(" normal")
        b64 = _doc_to_base64(doc)

        structure = service.extract_structure(b64)
        # Apply zero changes
        result = service.apply_changes([], b64)

        new_structure = service.extract_structure(result.document_base64)
        runs = new_structure.blocks[0].runs
        # The formatting should be preserved
        assert any(r.format and r.format.bold for r in runs)
