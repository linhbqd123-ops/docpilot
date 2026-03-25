"""Generate a sample .docx fixture for testing.

Run this script once to create tests/fixtures/sample.docx:
    python tests/create_fixture.py
"""

import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_sample_docx() -> Path:
    """Create a small .docx with headings, paragraphs, a list, and a table."""
    doc = Document()

    # Title
    title = doc.add_heading("Sample Document", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Heading 1
    doc.add_heading("Introduction", level=1)

    # Normal paragraph with mixed formatting
    p = doc.add_paragraph()
    run1 = p.add_run("This is a ")
    run2 = p.add_run("sample document")
    run2.bold = True
    run3 = p.add_run(" for testing DocPilot. It contains ")
    run4 = p.add_run("mixed formatting")
    run4.italic = True
    run5 = p.add_run(".")

    # Another paragraph
    doc.add_paragraph(
        "The document is used to verify that the DOC-MCP service "
        "can extract structure, apply changes, and preserve formatting."
    )

    # Heading 2
    doc.add_heading("Experience", level=2)

    # Bullet list
    for item in [
        "Developed AI-powered document tools",
        "Led a team of 5 engineers",
        "Reduced processing time by 40%",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # Heading 2
    doc.add_heading("Skills", level=2)

    # Table: 3x2
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    headers = ["Skill", "Level"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    data = [
        ("Python", "Expert"),
        ("TypeScript", "Advanced"),
    ]
    for r_idx, (skill, level) in enumerate(data, start=1):
        table.cell(r_idx, 0).text = skill
        table.cell(r_idx, 1).text = level

    # Save
    fixtures_dir = Path(__file__).parent / "fixtures"
    fixtures_dir.mkdir(parents=True, exist_ok=True)
    out_path = fixtures_dir / "sample.docx"
    doc.save(str(out_path))
    print(f"Created {out_path}")
    return out_path


if __name__ == "__main__":
    create_sample_docx()
